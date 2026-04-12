#!/usr/bin/env python3
"""
Generate the comprehensive system manual as a Word document (.docx)
with embedded Playwright screenshots.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOTS = Path(__file__).parent / "screenshots"
OUTPUT = Path(__file__).parent / "Manual_Sistema_Importacao_GrupoUnico.docx"


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_image_if_exists(doc, filename, width=Inches(6.5), caption=None):
    """Add screenshot image with optional caption."""
    img_path = SCREENSHOTS / filename
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        return True
    return False


def add_dual_theme(doc, base_name, fig_num, label, width=Inches(6)):
    """Add light + dark screenshots for a page."""
    add_image_if_exists(doc, f"{base_name}.png", width, f"Figura {fig_num}a — {label} (Tema Claro)")
    add_image_if_exists(doc, f"{base_name}_dark.png", width, f"Figura {fig_num}b — {label} (Tema Escuro)")


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "4338CA")  # Indigo

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, "F3F4F6")

    return table


def build_document():
    doc = Document()

    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ─── Styles ───
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)  # Indigo-900

    # ═══════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GRUPO UNI.CO")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Gestão de Importação")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual Completo do Sistema")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph()

    info_lines = [
        ("Versão:", "1.0"),
        ("Data:", "12 de Abril de 2026"),
        ("Classificação:", "Interno — Confidencial"),
        ("Departamento:", "Importação / Tecnologia"),
        ("URL Produção:", "http://192.168.168.124:8085"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + " ")
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════
    doc.add_heading("Sumário", level=1)

    toc_items = [
        "1. Visão Geral do Sistema",
        "2. Arquitetura e Stack Tecnológico",
        "3. Acesso e Autenticação",
        "4. Portal Principal",
        "5. Módulo de Importação",
        "   5.1  Dashboard",
        "   5.2  Dashboard Executivo",
        "   5.3  Meu Dia",
        "   5.4  Processos de Importação",
        "   5.5  Detalhe do Processo",
        "   5.6  Pré-Conferência",
        "   5.7  Câmbios",
        "   5.8  LIs / LPCOs",
        "   5.9  Desembaraço Aduaneiro",
        "   5.10 Numerário",
        "   5.11 Follow-Up",
        "   5.12 Comunicações",
        "   5.13 Alertas",
        "   5.14 Ingestão de E-mails",
        "   5.15 Auditoria",
        "   5.16 Configurações",
        "6. Módulo de Certificações",
        "   6.1  Dashboard",
        "   6.2  Validação",
        "   6.3  Produtos",
        "   6.4  Relatórios",
        "   6.5  Agendamentos",
        "   6.6  Configurações",
        "7. Fluxos de Negócio",
        "   7.1  Ciclo de Vida do Processo",
        "   7.2  Ingestão Automática de E-mails",
        "   7.3  Extração por IA",
        "   7.4  Validação Cruzada de Documentos",
        "   7.5  Geração de Espelho",
        "   7.6  Avanço Logístico Automático",
        "   7.7  Ciclo de Correções",
        "8. Integrações Externas",
        "9. Segurança",
        "10. Temas e Responsividade",
        "   10.1 Tema Claro e Escuro",
        "   10.2 Responsividade Mobile",
        "11. Glossário",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if not item.startswith("   "):
            for run in p.runs:
                run.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 1. VISÃO GERAL
    # ═══════════════════════════════════════════════════════
    doc.add_heading("1. Visão Geral do Sistema", level=1)
    doc.add_paragraph(
        "O Sistema de Gestão de Importação do Grupo Uni.co é uma plataforma web completa "
        "desenvolvida para automatizar e gerenciar todo o ciclo de vida dos processos de "
        "importação das marcas Puket e Imaginarium. O sistema integra inteligência artificial, "
        "automação de e-mails, validação cruzada de documentos e rastreamento logístico em "
        "uma única interface."
    )

    doc.add_heading("Principais Capacidades", level=2)
    capabilities = [
        "Gestão completa do ciclo de importação (11 estágios logísticos)",
        "Ingestão automática de e-mails com classificação de documentos por IA (Gemini 2.5 Flash)",
        "Extração inteligente de dados de invoices, packing lists, BLs e certificados",
        "28 verificações de validação cruzada entre documentos",
        "Geração automática de espelhos (planilhas de conferência) por marca",
        "Rastreamento de follow-up com 15+ marcos de acompanhamento",
        "Gestão de câmbio, LIs/LPCOs, numerários e desembaraço aduaneiro",
        "Integração com Google Drive, Sheets, Odoo, Gmail e Google Chat",
        "Módulo de certificações INMETRO/ANATEL com validação de estoque WMS/ERP",
        "Dashboard executivo com KPIs em tempo real",
        "Sistema completo de alertas e comunicações por e-mail",
        "Auditoria de ações e rastreamento de eventos por processo",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_heading("Módulos do Sistema", level=2)
    add_styled_table(doc,
        ["Módulo", "Descrição", "Acesso"],
        [
            ["Importação", "Gestão completa de processos de importação", "/importacao/*"],
            ["Certificações", "Validação INMETRO/ANATEL de produtos", "/certificacoes/*"],
            ["Portal", "Hub central com visão consolidada", "/portal"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 2. ARQUITETURA
    # ═══════════════════════════════════════════════════════
    doc.add_heading("2. Arquitetura e Stack Tecnológico", level=1)

    doc.add_heading("Visão da Arquitetura", level=2)
    doc.add_paragraph(
        "O sistema utiliza uma arquitetura monorepo com workspaces npm, composta por três "
        "aplicações: API (backend), Web (frontend) e Cert-API (microserviço de certificações). "
        "Todos os serviços são containerizados via Docker Compose."
    )

    doc.add_heading("Stack Frontend", level=2)
    add_styled_table(doc,
        ["Tecnologia", "Versão", "Propósito"],
        [
            ["React", "18", "Framework de UI"],
            ["Vite", "6", "Build tool e dev server"],
            ["Tailwind CSS", "4", "Framework CSS utility-first"],
            ["React Query", "—", "Cache e sincronização de dados"],
            ["React Hook Form + Zod", "—", "Formulários com validação"],
            ["React Router", "v7", "Roteamento SPA com lazy loading"],
            ["Recharts", "—", "Gráficos e visualizações"],
            ["Lucide Icons", "—", "Biblioteca de ícones"],
        ]
    )

    doc.add_heading("Stack Backend", level=2)
    add_styled_table(doc,
        ["Tecnologia", "Versão", "Propósito"],
        [
            ["Node.js", "20 LTS", "Runtime"],
            ["Express", "4", "Framework HTTP"],
            ["Drizzle ORM", "0.45+", "ORM com tipagem forte"],
            ["PostgreSQL", "16", "Banco de dados principal"],
            ["Redis", "—", "Cache (dashboard) e rate limiting"],
            ["Pino", "—", "Logging estruturado"],
            ["JWT", "—", "Autenticação stateless"],
            ["Multer", "2.1+", "Upload de arquivos (até 50MB)"],
            ["Nodemailer", "8.0+", "Envio de e-mails SMTP"],
            ["BullMQ", "—", "Filas de jobs assíncronos"],
        ]
    )

    doc.add_heading("Stack de IA", level=2)
    add_styled_table(doc,
        ["Tecnologia", "Propósito"],
        [
            ["Google AI Studio (Gemini 2.5 Flash)", "Extração de dados de documentos"],
            ["Modelo fallback: Gemini 2.0 Flash → 1.5 Pro", "Resiliência em caso de falha"],
            ["Temperature 0", "Respostas determinísticas"],
            ["Prompts especializados por tipo de documento", "Invoice, PL, BL, Draft BL, Certificado, E-mail"],
        ]
    )

    doc.add_heading("Infraestrutura", level=2)
    add_styled_table(doc,
        ["Componente", "Detalhe"],
        [
            ["Servidor", "192.168.168.124 (hostname: n8n)"],
            ["Docker Compose", "PostgreSQL + API + Web + Cert-API + Redis"],
            ["Nginx", "Reverse proxy com headers de segurança"],
            ["Portas Prod", "API: 127.0.0.1:3050, Web: 8085, DB: 127.0.0.1:5450"],
            ["Deploy", "bash scripts/deploy.sh (rsync + docker compose)"],
            ["CI/CD", "GitHub Actions (testes, lint, Trivy, build)"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 3. ACESSO E AUTENTICAÇÃO
    # ═══════════════════════════════════════════════════════
    doc.add_heading("3. Acesso e Autenticação", level=1)
    doc.add_paragraph(
        "O acesso ao sistema é realizado através de autenticação Google OAuth. "
        "Apenas usuários com e-mail do domínio autorizado e pertencentes ao grupo "
        "Google configurado podem acessar o sistema."
    )

    doc.add_heading("Tela de Login", level=2)
    add_image_if_exists(doc, "01_login.png", Inches(5.5), "Figura 1 — Tela de login com autenticação Google OAuth")

    doc.add_heading("Fluxo de Autenticação", level=2)
    auth_steps = [
        "O usuário acessa a URL do sistema e é redirecionado para a tela de login.",
        "Clica no botão 'Entrar com Google' e autentica com sua conta corporativa.",
        "O backend valida o token Google, verifica pertinência ao grupo autorizado.",
        "Um token JWT é gerado (validade 24h) e armazenado no navegador.",
        "O usuário é redirecionado ao Portal principal.",
    ]
    for i, step in enumerate(auth_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Perfis de Acesso", level=2)
    add_styled_table(doc,
        ["Perfil", "Permissões", "Exemplo"],
        [
            ["Admin", "Acesso total: CRUD de usuários, configurações, backfill, filas", "Nicolas, Administradores TI"],
            ["Analyst", "Operacional: processos, documentos, comunicações, follow-up", "Analistas de importação"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 4. PORTAL PRINCIPAL
    # ═══════════════════════════════════════════════════════
    doc.add_heading("4. Portal Principal", level=1)
    doc.add_paragraph(
        "O Portal é o ponto de entrada após o login. Apresenta uma visão consolidada "
        "dos dois módulos do sistema (Importação e Certificações) com resumos rápidos "
        "e acesso direto às funcionalidades mais utilizadas."
    )

    add_image_if_exists(doc, "02_portal.png", Inches(6), "Figura 2 — Portal principal com visão consolidada dos módulos")

    doc.add_heading("Elementos do Portal", level=2)
    portal_items = [
        "Saudação personalizada com nome do usuário e data atual",
        "Pills de resumo: processos ativos, vencidos, concluídos no mês",
        "Card do módulo Importação com estatísticas e link direto",
        "Card do módulo Certificações com estatísticas e link direto",
        "Links rápidos para ações frequentes",
        "Indicadores de saúde das APIs (Importação + Certificações)",
    ]
    for item in portal_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 5. MÓDULO DE IMPORTAÇÃO
    # ═══════════════════════════════════════════════════════
    doc.add_heading("5. Módulo de Importação", level=1)
    doc.add_paragraph(
        "O módulo de Importação é o core do sistema. Gerencia todo o ciclo de vida "
        "dos processos de importação, desde a recepção de documentos via e-mail até "
        "a internalização da mercadoria. A navegação é organizada em um sidebar lateral "
        "com seções temáticas."
    )

    add_image_if_exists(doc, "30_sidebar_importacao.png", Inches(2.5), "Figura 3 — Sidebar do módulo de Importação")

    # 5.1 Dashboard
    doc.add_heading("5.1 Dashboard", level=2)
    doc.add_paragraph(
        "O Dashboard apresenta uma visão geral operacional com KPIs, gráficos e "
        "tabelas de acompanhamento rápido. Os dados são cacheados no Redis com "
        "TTL de 60 segundos para performance."
    )
    add_dual_theme(doc, "03_dashboard", 4, "Dashboard operacional com KPIs e gráficos")

    doc.add_heading("Componentes do Dashboard", level=3)
    dash_items = [
        "KPI Cards: Processos ativos, vencidos, concluídos, valor FOB total",
        "Gráfico de barras: Processos por status",
        "Gráfico de linha: Tendência mensal de processos",
        "Gráfico de pizza: Distribuição de FOB por marca",
        "Tabela de alertas recentes (últimos 10)",
        "Tabela de processos recentes com status",
        "Logs de e-mail do dia",
        "Painel de SLA com indicadores de cumprimento",
    ]
    for item in dash_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 5.2 Executivo
    doc.add_heading("5.2 Dashboard Executivo", level=2)
    doc.add_paragraph(
        "Visão de alto nível para gestores com métricas estratégicas e "
        "indicadores de performance consolidados."
    )
    add_dual_theme(doc, "04_executivo", 5, "Dashboard executivo")

    doc.add_page_break()

    # 5.3 Meu Dia
    doc.add_heading("5.3 Meu Dia", level=2)
    doc.add_paragraph(
        "Painel pessoal que concentra as tarefas e prioridades do dia do "
        "analista logado, facilitando o foco nas atividades mais urgentes."
    )
    add_image_if_exists(doc, "19_meu_dia.png", Inches(6), "Figura 6 — Tela Meu Dia")

    doc.add_page_break()

    # 5.4 Processos
    doc.add_heading("5.4 Processos de Importação", level=2)
    doc.add_paragraph(
        "A tela de processos é o coração operacional do sistema. Lista todos os processos "
        "de importação com filtros avançados, busca por texto, paginação e acesso rápido "
        "aos detalhes de cada processo."
    )
    add_dual_theme(doc, "05_processos_lista", 7, "Lista de processos com filtros")

    doc.add_heading("Filtros Disponíveis", level=3)
    add_styled_table(doc,
        ["Filtro", "Descrição"],
        [
            ["Busca por texto", "Pesquisa por código do processo, fornecedor ou referência"],
            ["Status", "Draft, Docs Recebidos, Validando, Validado, Espelho, Fenicia, LI, Concluído"],
            ["Marca", "Puket ou Imaginarium"],
            ["Período", "Data inicial e final de criação"],
        ]
    )

    doc.add_heading("Criar Novo Processo", level=3)
    doc.add_paragraph(
        "Novos processos podem ser criados manualmente ou a partir de itens da Pré-Conferência. "
        "O formulário solicita código do processo, fornecedor, marca, incoterm e dados básicos."
    )
    add_image_if_exists(doc, "07_processo_novo.png", Inches(6), "Figura 8 — Formulário de criação de novo processo")

    doc.add_page_break()

    # 5.5 Detalhe do Processo
    doc.add_heading("5.5 Detalhe do Processo", level=2)
    doc.add_paragraph(
        "A tela de detalhe apresenta todas as informações de um processo organizado "
        "em abas temáticas. Inclui informações gerais, timeline de eventos, "
        "estágio logístico atual e sub-módulos específicos."
    )
    add_dual_theme(doc, "06_processo_detalhe", 9, "Visão geral do detalhe do processo")

    doc.add_heading("Abas do Processo", level=3)

    # Draft BL
    doc.add_heading("Draft BL", level=3)
    doc.add_paragraph(
        "Gerenciamento do Draft Bill of Lading com upload, checklist de 10 itens, "
        "extração por IA e comparação com o BL final (Revisado)."
    )
    add_image_if_exists(doc, "06a_processo_draft_bl.png", Inches(6), "Figura 10 — Aba Draft BL")

    # Documentos
    doc.add_heading("Documentos", level=3)
    doc.add_paragraph(
        "Lista de todos os documentos associados ao processo com upload, visualização, "
        "classificação automática e dados extraídos por IA."
    )
    add_image_if_exists(doc, "06b_processo_documentos.png", Inches(6), "Figura 11 — Aba de Documentos")

    # Validação
    doc.add_heading("Validação", level=3)
    doc.add_paragraph(
        "Resultado das 28 verificações de validação cruzada entre documentos. "
        "Mostra status (aprovado/reprovado/alerta), valores esperados vs. encontrados "
        "e documentos comparados."
    )
    add_image_if_exists(doc, "06c_processo_validacao.png", Inches(6), "Figura 12 — Aba de Validação")

    # Follow-Up
    doc.add_heading("Follow-Up", level=3)
    doc.add_paragraph(
        "Timeline de marcos do processo com 15+ checkpoints rastreados. "
        "Calcula progresso percentual e sincroniza com a planilha Follow-Up no Google Sheets."
    )
    add_image_if_exists(doc, "06d_processo_follow_up.png", Inches(6), "Figura 13 — Aba de Follow-Up")

    # Comunicações
    doc.add_heading("Comunicações", level=3)
    doc.add_paragraph(
        "Histórico de e-mails enviados e recebidos do processo, com drafts, "
        "templates e assinaturas personalizáveis."
    )
    add_image_if_exists(doc, "06e_processo_comunicacoes.png", Inches(6), "Figura 14 — Aba de Comunicações")

    # Histórico
    doc.add_heading("Histórico / Timeline", level=3)
    doc.add_paragraph(
        "Timeline completa de todos os eventos do processo: criação, uploads, "
        "transições de status, validações, envios e marcos logísticos."
    )
    add_image_if_exists(doc, "06f_processo_historico.png", Inches(6), "Figura 15 — Aba de Histórico")

    doc.add_page_break()

    # 5.6 Pré-Conferência
    doc.add_heading("5.6 Pré-Conferência", level=2)
    doc.add_paragraph(
        "Módulo de sincronização com as planilhas KIOM de pré-conferência. "
        "Permite importar itens, visualizar pedidos pendentes e criar processos "
        "diretamente a partir dos dados pré-conferidos."
    )
    add_image_if_exists(doc, "08_pre_conferencia.png", Inches(6), "Figura 16 — Tela de Pré-Conferência")

    doc.add_heading("Funcionalidades", level=3)
    precons_items = [
        "Sincronização de planilhas KIOM (KIOM IMAG, KIOM PUK, etc.)",
        "Parsing automático de colunas: Pedido, ETD, Coleta, Fornecedor, etc.",
        "Criação de processos a partir de itens da pré-conferência",
        "Visualização e filtro de itens pendentes",
    ]
    for item in precons_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 5.7 Câmbios
    doc.add_heading("5.7 Câmbios", level=2)
    doc.add_paragraph(
        "Gerenciamento de operações de câmbio (USD → BRL) associadas aos processos. "
        "Controla taxa, valores, prazos de pagamento e vencimentos."
    )
    add_image_if_exists(doc, "09_cambios.png", Inches(6), "Figura 17 — Tela de Câmbios")

    add_styled_table(doc,
        ["Campo", "Descrição"],
        [
            ["Tipo", "Saldo (balance) ou Depósito (deposit)"],
            ["Valor USD", "Montante em dólares americanos"],
            ["Taxa de Câmbio", "Cotação USD/BRL"],
            ["Valor BRL", "Montante convertido em reais"],
            ["Prazo Pagamento", "Data limite para efetuar o pagamento"],
            ["Data Vencimento", "Data de expiração da operação"],
        ]
    )

    doc.add_page_break()

    # 5.8 LIs / LPCOs
    doc.add_heading("5.8 LIs / LPCOs", level=2)
    doc.add_paragraph(
        "Rastreamento de Licenças de Importação (LI) e LPCOs. Controla prazos de "
        "deferimento, números LPCO, validade e status de cada item que requer licença. "
        "Itens com NCM iniciando em 39, 42, 61-65, 85 ou 95 são automaticamente "
        "identificados como requerendo LI."
    )
    add_image_if_exists(doc, "10_lis_lpcos.png", Inches(6), "Figura 18 — Tela de LIs / LPCOs")

    doc.add_page_break()

    # 5.9 Desembaraço
    doc.add_heading("5.9 Desembaraço Aduaneiro", level=2)
    doc.add_paragraph(
        "Acompanhamento do processo de desembaraço aduaneiro com rastreamento "
        "de status, documentação e prazos."
    )
    add_image_if_exists(doc, "11_desembaraco.png", Inches(6), "Figura 19 — Tela de Desembaraço Aduaneiro")

    doc.add_page_break()

    # 5.10 Numerário
    doc.add_heading("5.10 Numerário", level=2)
    doc.add_paragraph(
        "Gestão de numerários com controle de pagamentos, depósitos e saldos "
        "relacionados aos processos de importação."
    )
    add_image_if_exists(doc, "12_numerario.png", Inches(6), "Figura 20 — Tela de Numerário")

    doc.add_page_break()

    # 5.11 Follow-Up
    doc.add_heading("5.11 Follow-Up", level=2)
    doc.add_paragraph(
        "Visão consolidada de todos os marcos de acompanhamento dos processos. "
        "Permite visualizar em grade ou calendário, com indicadores de status "
        "(no prazo, próximo do vencimento, vencido)."
    )
    add_image_if_exists(doc, "13_follow_up.png", Inches(6), "Figura 21 — Tela de Follow-Up")

    doc.add_heading("Marcos de Acompanhamento (15+ checkpoints)", level=3)
    add_styled_table(doc,
        ["#", "Marco", "Descrição"],
        [
            ["1", "documentsReceivedAt", "Data de recebimento dos documentos"],
            ["2", "preInspectionAt", "Pré-inspeção (pós-validação)"],
            ["3", "savedToFolderAt", "Documentos salvos no Drive"],
            ["4", "ncmVerifiedAt", "NCM verificado"],
            ["5", "ncmBlCheckedAt", "NCM vs BL conferido"],
            ["6", "freightBlCheckedAt", "Frete do BL conferido"],
            ["7", "espelhoBuiltAt", "Espelho construído"],
            ["8", "espelhoGeneratedAt", "Espelho gerado (XLSX)"],
            ["9", "invoiceSentFeniciaAt", "Invoice enviada à Fenicia"],
            ["10", "signaturesCollectedAt", "Assinaturas coletadas"],
            ["11", "signedDocsSentAt", "Docs assinados enviados"],
            ["12", "sentToFeniciaAt", "Documentos enviados à Fenicia"],
            ["13", "diDraftAt", "Minuta da DI"],
            ["14", "liSubmittedAt", "LI submetida"],
            ["15", "liApprovedAt", "LI aprovada"],
            ["16", "liDeadline", "Prazo da LI (embarque + 13 dias)"],
        ]
    )

    doc.add_page_break()

    # 5.12 Comunicações
    doc.add_heading("5.12 Comunicações", level=2)
    doc.add_paragraph(
        "Central de comunicações por e-mail com templates pré-configurados, "
        "assinaturas personalizáveis (até 4 por usuário) e histórico completo. "
        "Envio via SMTP (mta.imgnet.com.br:2525) com TLS obrigatório em produção."
    )
    add_image_if_exists(doc, "14_comunicacoes.png", Inches(6), "Figura 22 — Tela de Comunicações")

    doc.add_heading("Templates Disponíveis", level=3)
    add_styled_table(doc,
        ["Template", "Descrição", "Uso"],
        [
            ["Correção KIOM", "Solicita correção de documentos ao KIOM", "Quando validação falha"],
            ["Submissão Fenicia", "Envia documentos + espelho à Fenicia", "Pós-validação aprovada"],
            ["Certificação ISA", "Fluxo de certificação ISA", "Processos com certificação"],
        ]
    )

    doc.add_page_break()

    # 5.13 Alertas
    doc.add_heading("5.13 Alertas", level=2)
    doc.add_paragraph(
        "Central de alertas com três níveis de severidade (info, warning, critical). "
        "Alertas críticos e warnings são automaticamente enviados ao Google Chat via webhook. "
        "Alertas são gerados automaticamente por: falhas de validação, prazos excedidos, "
        "processos estagnados e falhas na ingestão de e-mails."
    )
    add_image_if_exists(doc, "15_alertas.png", Inches(6), "Figura 23 — Central de Alertas")

    doc.add_page_break()

    # 5.14 Email Ingestion
    doc.add_heading("5.14 Ingestão de E-mails", level=2)
    doc.add_paragraph(
        "Painel de monitoramento da ingestão automática de e-mails. O sistema verifica "
        "a caixa de entrada (eduarda.souza@grupounico.com) via Gmail API a cada 5 minutos, "
        "classificando anexos e associando a processos automaticamente."
    )
    add_image_if_exists(doc, "16_email_ingestion.png", Inches(6), "Figura 24 — Painel de Ingestão de E-mails")

    doc.add_heading("Pipeline de Ingestão", level=3)
    ingestion_steps = [
        "Verificação de novos e-mails via Gmail API (a cada 5 minutos)",
        "Extração do código do processo via regex no assunto/corpo",
        "Classificação de anexos por nome de arquivo + contexto IA",
        "Extração de texto (PDF, Excel, Word)",
        "Chamada de IA para extração estruturada de dados",
        "Upload para Google Drive na pasta do processo",
        "Inserção no banco de dados com dados extraídos",
        "Atualização automática de status se todos os 3 docs principais recebidos",
        "Registro de marco no follow-up",
    ]
    for i, step in enumerate(ingestion_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # 5.15 Auditoria
    doc.add_heading("5.15 Auditoria", level=2)
    doc.add_paragraph(
        "Log de auditoria com registro completo de todas as ações realizadas no sistema. "
        "Inclui usuário, ação, entidade afetada, detalhes e endereço IP."
    )
    add_image_if_exists(doc, "17_auditoria.png", Inches(6), "Figura 25 — Log de Auditoria")

    doc.add_page_break()

    # 5.16 Configurações
    doc.add_heading("5.16 Configurações", level=2)
    doc.add_paragraph(
        "Painel de configurações do sistema, incluindo gestão de usuários, "
        "assinaturas de e-mail, integrações e parâmetros gerais. "
        "Acesso restrito a administradores."
    )
    add_image_if_exists(doc, "18_configuracoes.png", Inches(6), "Figura 26 — Tela de Configurações")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 6. MÓDULO DE CERTIFICAÇÕES
    # ═══════════════════════════════════════════════════════
    doc.add_heading("6. Módulo de Certificações", level=1)
    doc.add_paragraph(
        "O módulo de Certificações gerencia a validação de conformidade INMETRO e ANATEL "
        "dos produtos importados. Integra com o WMS Oracle (estoque físico) e ERPs SQL Server "
        "(Puket e Imaginarium) para validação de estoque e certificação."
    )

    add_image_if_exists(doc, "31_sidebar_certificacoes.png", Inches(2.5), "Figura 27 — Sidebar do módulo de Certificações")

    # 6.1 Dashboard
    doc.add_heading("6.1 Dashboard de Certificações", level=2)
    doc.add_paragraph(
        "Visão geral do módulo com estatísticas de produtos, status de validação "
        "e resumo de relatórios."
    )
    add_image_if_exists(doc, "20_cert_dashboard.png", Inches(6), "Figura 28 — Dashboard de Certificações")

    doc.add_page_break()

    # 6.2 Validação
    doc.add_heading("6.2 Validação", level=2)
    doc.add_paragraph(
        "Execução de validação de certificações com verificação automática de "
        "estoque no WMS e conformidade nos ERPs."
    )
    add_image_if_exists(doc, "21_cert_validacao.png", Inches(6), "Figura 29 — Tela de Validação de Certificações")

    doc.add_page_break()

    # 6.3 Produtos
    doc.add_heading("6.3 Produtos", level=2)
    doc.add_paragraph(
        "Listagem e gestão de produtos certificados com informações de SKU, "
        "status de certificação e dados de estoque."
    )
    add_image_if_exists(doc, "22_cert_produtos.png", Inches(6), "Figura 30 — Lista de Produtos")

    doc.add_page_break()

    # 6.4 Relatórios
    doc.add_heading("6.4 Relatórios", level=2)
    doc.add_paragraph(
        "Relatórios gerados pelas validações com detalhamento dos resultados, "
        "produtos conformes e não-conformes."
    )
    add_image_if_exists(doc, "23_cert_relatorios.png", Inches(6), "Figura 31 — Relatórios de Certificação")

    doc.add_page_break()

    # 6.5 Agendamentos
    doc.add_heading("6.5 Agendamentos", level=2)
    doc.add_paragraph(
        "Configuração de agendamentos para execução automática de validações "
        "em intervalos definidos."
    )
    add_image_if_exists(doc, "24_cert_agendamentos.png", Inches(6), "Figura 32 — Agendamentos de Certificação")

    doc.add_page_break()

    # 6.6 Configurações
    doc.add_heading("6.6 Configurações de Certificações", level=2)
    doc.add_paragraph(
        "Parâmetros específicos do módulo de certificações: fontes de dados, "
        "credenciais de acesso a WMS/ERP e configurações de validação."
    )
    add_image_if_exists(doc, "25_cert_configuracoes.png", Inches(6), "Figura 33 — Configurações de Certificações")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 7. FLUXOS DE NEGÓCIO
    # ═══════════════════════════════════════════════════════
    doc.add_heading("7. Fluxos de Negócio", level=1)

    # 7.1 Ciclo de Vida
    doc.add_heading("7.1 Ciclo de Vida do Processo de Importação", level=2)
    doc.add_paragraph(
        "Cada processo de importação segue um ciclo de vida controlado por uma "
        "máquina de estados que garante transições válidas. O sistema impede "
        "transições inválidas e registra cada mudança na timeline."
    )

    doc.add_heading("Status do Processo", level=3)
    add_styled_table(doc,
        ["Status", "Descrição", "Transição Seguinte"],
        [
            ["draft", "Rascunho inicial", "documents_received"],
            ["documents_received", "Documentos recebidos (3 principais)", "validating"],
            ["validating", "Validação em execução", "validated / pending_correction"],
            ["validated", "Todos os 28 checks aprovados", "espelho_generated"],
            ["pending_correction", "Correções necessárias", "validating (re-validar)"],
            ["espelho_generated", "Espelho XLSX gerado", "sent_to_fenicia"],
            ["sent_to_fenicia", "Docs enviados à Fenicia", "li_pending"],
            ["li_pending", "Aguardando aprovação da LI", "completed"],
            ["completed", "Processo concluído", "—"],
            ["cancelled", "Processo cancelado", "—"],
        ]
    )

    doc.add_heading("Estágios Logísticos (11 estágios)", level=3)
    doc.add_paragraph(
        "Além do status do processo, existe um rastreamento logístico paralelo "
        "com 11 estágios derivados automaticamente dos marcos do follow-up:"
    )
    add_styled_table(doc,
        ["#", "Estágio", "Condição de Avanço"],
        [
            ["1", "consolidation", "Padrão inicial — documentos em consolidação"],
            ["2", "booked", "Reserva de embarque confirmada"],
            ["3", "shipped", "Mercadoria embarcada (shipmentDate preenchido)"],
            ["4", "in_transit", "Em trânsito marítimo/aéreo"],
            ["5", "arrived_port", "Chegou ao porto de destino"],
            ["6", "customs_started", "Desembaraço iniciado"],
            ["7", "customs_cleared", "Desembaraço concluído"],
            ["8", "in_transit_domestic", "Em trânsito doméstico"],
            ["9", "delivered_warehouse", "Entregue no armazém"],
            ["10", "inspected", "Inspeção física concluída"],
            ["11", "internalized", "Mercadoria internalizada"],
        ]
    )

    doc.add_page_break()

    # 7.2 Ingestão de E-mails
    doc.add_heading("7.2 Fluxo de Ingestão Automática de E-mails", level=2)
    doc.add_paragraph(
        "O sistema verifica a caixa de e-mail a cada 5 minutos via Gmail API. "
        "E-mails são processados, classificados e seus anexos são automaticamente "
        "associados aos processos correspondentes."
    )

    flow_steps = [
        ("Verificação", "Gmail API busca novos e-mails (cron a cada 5 min)"),
        ("Extração de Código", "Regex identifica código do processo no assunto/corpo (IMP-xxxx, PUKxxx, etc.)"),
        ("Classificação", "Anexos classificados por nome de arquivo: invoice → Invoice, pl/packing → Packing List, bl → Bill of Lading"),
        ("Análise IA", "Gemini 2.5 Flash analisa contexto do e-mail para refinar classificação"),
        ("Extração de Texto", "PDF parsing (pdf-parse), Excel (xlsx), Word para texto plano"),
        ("Extração IA", "IA extrai dados estruturados: fornecedor, itens, valores, datas, pesos"),
        ("Upload Drive", "Arquivo enviado ao Google Drive na pasta do processo com nome padronizado"),
        ("Registro no BD", "Documento inserido no banco com metadados e dados extraídos (aiParsedData JSONB)"),
        ("Auto-Transição", "Se invoice + packing list + BL presentes → status = documents_received"),
        ("Sincronização", "Marco documentsReceivedAt atualizado no follow-up + sync Google Sheets"),
    ]
    add_styled_table(doc,
        ["Etapa", "Descrição"],
        [[s, d] for s, d in flow_steps]
    )

    doc.add_page_break()

    # 7.3 IA
    doc.add_heading("7.3 Extração por Inteligência Artificial", level=2)
    doc.add_paragraph(
        "O sistema utiliza o modelo Gemini 2.5 Flash (via Google AI Studio) para "
        "extrair dados estruturados de documentos comerciais. Cada tipo de documento "
        "possui um prompt especializado e schema de validação Zod."
    )

    doc.add_heading("Tipos de Extração", level=3)
    add_styled_table(doc,
        ["Tipo", "Campos Extraídos"],
        [
            ["Invoice", "Fornecedor, nº fatura, data, itens (código, descrição, qtd, preço unit., total), FOB total, termos de pagamento, incoterm"],
            ["Packing List", "Itens, quantidades, peso líquido/bruto por item, nº de caixas, CBM, dimensões"],
            ["Bill of Lading", "Embarcador, consignatário, navio, nº BL, porto embarque/descarga, cia marítima, ETD, ETA, tipo container"],
            ["Draft BL", "Mesmo que BL, marcado como draft para comparação posterior"],
            ["Certificado", "Tipo (fitossanitário, COA, INMETRO, etc.), número, emissor, validade"],
            ["E-mail", "Código do processo, tipos de documentos, urgência, categoria, datas-chave"],
        ]
    )

    doc.add_heading("Governança de IA", level=3)
    ai_gov = [
        "Todas as requisições à IA são logadas com modelo, tokens e latência",
        "Scores de confiança por campo (0-1) para identificar extrações duvidosas",
        "Fallback automático: Gemini 2.5 Flash → 2.0 Flash → 1.5 Pro",
        "Temperature 0 para resultados determinísticos",
        "Sanitização de prefixos espúrios (layout bleed de PDFs)",
        "Timeout de 90 segundos por requisição",
    ]
    for item in ai_gov:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 7.4 Validação
    doc.add_heading("7.4 Validação Cruzada de Documentos", level=2)
    doc.add_paragraph(
        "O motor de validação executa 28 verificações independentes em paralelo, "
        "comparando dados entre Invoice, Packing List, Bill of Lading e dados do "
        "processo/follow-up. Cada check retorna status (passed/failed/warning), "
        "valor esperado vs. encontrado e documentos comparados."
    )

    doc.add_heading("Lista Completa de Verificações (28)", level=3)
    checks = [
        ["1", "Box Quantity Match", "Qtd de caixas: PL vs Invoice"],
        ["2", "CBM Match", "CBM: PL vs Follow-Up"],
        ["3", "Container Type vs FUP", "Tipo de container vs Follow-Up"],
        ["4", "Currency Check", "Moeda USD/BRL consistente"],
        ["5", "Date Sequence", "ETD ≤ ETA (sequência válida)"],
        ["6", "Dates Match", "Datas entre Invoice, PL e BL"],
        ["7", "Description vs Odoo", "Descrição dos itens vs catálogo Odoo"],
        ["8", "Exporter Match", "Exportador: Invoice vs BL"],
        ["9", "FOB Calculation", "Soma dos itens = total FOB"],
        ["10", "Freight Value Match", "Frete: Invoice vs BL"],
        ["11", "Freight vs FUP", "Frete vs Follow-Up"],
        ["12", "Gross Weight Match", "Peso bruto: PL vs Follow-Up"],
        ["13", "Importer Match", "Importador: Invoice vs BL"],
        ["14", "Incoterm Check", "Incoterm consistente entre docs"],
        ["15", "Invoice Value vs FUP", "Valor da invoice vs Follow-Up"],
        ["16", "Item-Level Match", "Match detalhado de SKUs entre docs"],
        ["17", "Manufacturer Completeness", "Completude dos dados do fabricante"],
        ["18", "NCM-BL Description", "NCM vs descrição no BL"],
        ["19", "Net Weight Match", "Peso líquido entre docs"],
        ["20", "Payment Terms", "Termos de pagamento consistentes"],
        ["21", "Ports Match", "Portos de embarque/descarga"],
        ["22", "Process Reference", "Referência do processo nos docs"],
        ["23", "Supplier Address", "Endereço do fornecedor entre docs"],
        ["24", "Unit Type Validation", "Unidades de medida válidas"],
        ["25", "Weight Ratio", "Proporção peso bruto/líquido coerente"],
        ["26", "Exporter Name/Address", "Nome/endereço exportador"],
        ["27", "Importer Name/Address", "Nome/endereço importador"],
        ["28", "CBM vs FUP", "CBM vs Follow-Up (recálculo)"],
    ]
    add_styled_table(doc, ["#", "Verificação", "Descrição"], checks)

    doc.add_heading("Resultado da Validação", level=3)
    doc.add_paragraph(
        "Se TODOS os 28 checks passam → status = validated e marco preInspectionAt atualizado. "
        "Se QUALQUER check falha → status = pending_correction, documentos movidos para "
        "pasta de correção no Drive e e-mail de correção KIOM gerado automaticamente como draft."
    )

    doc.add_page_break()

    # 7.5 Espelho
    doc.add_heading("7.5 Geração de Espelho", level=2)
    doc.add_paragraph(
        "O Espelho é uma planilha XLSX de conferência gerada automaticamente "
        "após a validação, utilizando templates específicos por marca (Puket/Imaginarium). "
        "Identifica automaticamente itens que requerem LI baseado no prefixo NCM."
    )

    espelho_steps = [
        "Busca processo + itens validados",
        "Identifica itens com LI obrigatória (NCM 39, 42, 61-65, 85, 95)",
        "Aplica template da marca (Puket ou Imaginarium)",
        "Gera arquivo XLSX com cálculos e formatação",
        "Upload ao Google Drive + registro no banco (com versionamento: v1, v2...)",
        "Status → espelho_generated, marco espelhoGeneratedAt atualizado",
        "Sincronização com Follow-Up Google Sheet",
    ]
    for i, step in enumerate(espelho_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # 7.6 Logístico
    doc.add_heading("7.6 Avanço Logístico Automático", level=2)
    doc.add_paragraph(
        "O sistema executa a cada 30 minutos (cron) uma derivação automática do "
        "estágio logístico de cada processo, baseado nos marcos preenchidos no follow-up. "
        "Isso elimina a necessidade de atualização manual do tracking logístico."
    )

    doc.add_heading("Regras de Derivação", level=3)
    doc.add_paragraph(
        "O estágio é derivado em cascata: o sistema verifica de trás para frente "
        "qual o marco mais avançado preenchido e define o estágio correspondente. "
        "Também calcula automaticamente o prazo da LI (data de embarque + 13 dias) "
        "e gera alertas se o prazo for excedido."
    )

    doc.add_page_break()

    # 7.7 Correções
    doc.add_heading("7.7 Ciclo de Correções", level=2)
    doc.add_paragraph(
        "Quando a validação identifica inconsistências, o sistema inicia um ciclo "
        "de correção automatizado:"
    )

    correction_steps = [
        "Validação falha → status = pending_correction",
        "Documentos movidos para pasta /correção/ no Drive",
        "E-mail de correção KIOM gerado como draft com detalhes dos erros",
        "Analista revisa e envia o e-mail ao KIOM",
        "KIOM responde com documentos corrigidos",
        "Ingestão de e-mails captura os novos documentos",
        "Novos documentos são processados e versões anteriores atualizadas",
        "Analista executa re-validação",
        "Se aprovado: documentos movidos de volta, status = validated",
        "Marco preInspectionAt atualizado e ciclo encerrado",
    ]
    for i, step in enumerate(correction_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 8. INTEGRAÇÕES
    # ═══════════════════════════════════════════════════════
    doc.add_heading("8. Integrações Externas", level=1)

    add_styled_table(doc,
        ["Serviço", "Protocolo", "Uso"],
        [
            ["Google Drive", "REST API (OAuth 2.0)", "Armazenamento de documentos, organização por processo em pastas"],
            ["Google Sheets", "REST API (OAuth 2.0)", "Sincronização de marcos Follow-Up em planilha compartilhada"],
            ["Gmail", "REST API (OAuth 2.0)", "Ingestão automática de e-mails (eduarda.souza@grupounico.com)"],
            ["Google Chat", "Webhook", "Notificações de alertas críticos e warnings"],
            ["Google AI Studio", "OpenAI-compatible API", "Extração IA com Gemini 2.5 Flash"],
            ["Odoo", "XML-RPC", "Consulta de catálogo de produtos para validação"],
            ["SMTP (mta.imgnet.com.br)", "SMTP/TLS (porta 2525)", "Envio de e-mails (from: global@grupounico.com)"],
            ["WMS Oracle", "Oracle DB (1521)", "Estoque físico para certificações (192.168.168.10)"],
            ["ERP Puket (SQL Server)", "TDS", "Dados de produtos Puket (db01.grupounico.com)"],
            ["ERP Imaginarium (SQL Server)", "TDS", "Dados de produtos Imaginarium (db02.grupounico.com)"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 9. SEGURANÇA
    # ═══════════════════════════════════════════════════════
    doc.add_heading("9. Segurança", level=1)
    doc.add_paragraph(
        "O sistema implementa múltiplas camadas de segurança, validadas por "
        "pentest em 17 categorias (abril/2026). Os principais controles:"
    )

    add_styled_table(doc,
        ["Camada", "Controle", "Detalhe"],
        [
            ["Autenticação", "JWT + Google OAuth", "Token 24h, validação de grupo Google"],
            ["Autorização", "RBAC (admin/analyst)", "Middleware adminMiddleware para rotas restritas"],
            ["Rede", "Bind 127.0.0.1", "PostgreSQL e API não expostos externamente"],
            ["Headers", "Nginx security headers", "CSP, X-Frame-Options, HSTS, server_tokens off"],
            ["Rate Limiting", "Redis-backed", "INCR+EXPIRE com X-RateLimit-* headers"],
            ["XSS", "DOMPurify", "Sanitização em todo dangerouslySetInnerHTML"],
            ["SMTP", "TLS + sanitização CRLF", "rejectUnauthorized=true em produção"],
            ["Upload", "Multer + validação", "Limite 50MB, tipos permitidos, rate limit 20/min"],
            ["Senhas", "bcryptjs (10 rounds)", "Mínimo 8 caracteres"],
            ["Docker", "Variáveis obrigatórias", "${VAR:?required} para credenciais"],
            ["Timeouts", "30-90s por serviço", "Drive/Gmail/Sheets: 30s, Odoo: 30s, AI: 90s"],
            ["Auditoria", "Log completo", "Todas as ações logadas com IP e detalhes"],
            ["JSON", "Limite 2MB", "body-parser com limite de payload"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 10. TEMAS E RESPONSIVIDADE
    # ═══════════════════════════════════════════════════════
    doc.add_heading("10. Temas e Responsividade", level=1)

    doc.add_heading("10.1 Tema Claro e Escuro", level=2)
    doc.add_paragraph(
        "O sistema suporta dois temas visuais: Claro (padrão) e Escuro. "
        "O tema pode ser alternado através do botão de configuração no header "
        "(ícone sol/lua), com três opções: Claro, Escuro e Sistema "
        "(segue a preferência do sistema operacional). A preferência é "
        "salva no navegador e mantida entre sessões."
    )

    doc.add_heading("Comparativo Visual — Portal", level=3)
    add_dual_theme(doc, "02_portal", 40, "Portal")
    doc.add_heading("Comparativo Visual — Meu Dia", level=3)
    add_dual_theme(doc, "19_meu_dia", 41, "Meu Dia")
    doc.add_heading("Comparativo Visual — Certificações", level=3)
    add_dual_theme(doc, "20_cert_dashboard", 42, "Dashboard Certificações")
    doc.add_heading("Comparativo Visual — Configurações", level=3)
    add_dual_theme(doc, "18_configuracoes", 43, "Configurações")

    doc.add_page_break()

    doc.add_heading("10.2 Responsividade Mobile", level=2)
    doc.add_paragraph(
        "Todas as telas do sistema são responsivas e adaptam-se automaticamente "
        "a diferentes tamanhos de tela (desktop, tablet e smartphone). "
        "Em dispositivos móveis, o sidebar colapsa em um menu hamburger, "
        "KPIs empilham verticalmente, tabelas ganham scroll horizontal e "
        "formulários reorganizam-se em coluna única."
    )

    doc.add_heading("Exemplos Mobile (iPhone 14 — 390x844)", level=3)
    mobile_examples = [
        ("01_login_mobile", "Login Mobile"),
        ("02_portal_mobile", "Portal Mobile"),
        ("03_dashboard_mobile", "Dashboard Mobile"),
        ("04_executivo_mobile", "Executivo Mobile"),
        ("05_processos_mobile", "Processos Mobile"),
        ("06_processo_detalhe_mobile", "Detalhe do Processo Mobile"),
        ("19_meu_dia_mobile", "Meu Dia Mobile"),
        ("20_cert_dashboard_mobile", "Certificações Mobile"),
    ]
    fig = 50
    for fname, label in mobile_examples:
        add_image_if_exists(doc, f"{fname}.png", Inches(2.5), f"Figura {fig} — {label}")
        fig += 1

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 11. GLOSSÁRIO
    # ═══════════════════════════════════════════════════════
    doc.add_heading("11. Glossário", level=1)

    glossary = [
        ["BL (Bill of Lading)", "Conhecimento de embarque — documento que comprova o transporte marítimo"],
        ["Draft BL", "Versão prévia do BL para conferência antes da emissão final"],
        ["CBM", "Cubic Meter — metro cúbico, unidade de volume para carga"],
        ["DI", "Declaração de Importação — documento aduaneiro"],
        ["Espelho", "Planilha de conferência gerada pelo sistema com dados consolidados"],
        ["ETD / ETA", "Estimated Time of Departure / Arrival — previsão de embarque/chegada"],
        ["Fenicia", "Despachante aduaneiro — recebe documentos para desembaraço"],
        ["FOB", "Free on Board — valor da mercadoria no porto de origem"],
        ["Incoterm", "International Commercial Terms — termos de negociação (FOB, CIF, etc.)"],
        ["KIOM", "Trading/agente de compras — responsável pelos documentos de exportação"],
        ["LI / LPCO", "Licença de Importação / Licença, Permissão, Certificado e Outros"],
        ["NCM", "Nomenclatura Comum do Mercosul — código de classificação fiscal"],
        ["PL (Packing List)", "Lista de embarque — detalha conteúdo de cada volume"],
        ["Pré-Conferência", "Dados recebidos antes do embarque para planejamento"],
        ["WMS", "Warehouse Management System — sistema de gestão de armazém"],
    ]
    add_styled_table(doc, ["Termo", "Definição"], glossary)

    # ═══════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento gerado automaticamente pelo Sistema de Gestão de Importação")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Grupo Uni.co — Departamento de Importação / Tecnologia")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abril 2026 — Versão 1.0")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Save
    doc.save(str(OUTPUT))
    print(f"Manual saved to: {OUTPUT}")
    print(f"Pages estimated: ~{len(doc.paragraphs) // 8}")


if __name__ == "__main__":
    build_document()
